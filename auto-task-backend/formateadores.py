import re
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import xml.etree.ElementTree as ET

# Intentar importar latex2mathml
try:
    from latex2mathml.converter import convert as latex2mathml
    LATEX_SUPPORT = True
except ImportError:
    LATEX_SUPPORT = False
    print("⚠️ latex2mathml no instalado. Ejecuta: pip install latex2mathml")

def add_math_equation(paragraph, latex_code):
    """
    Inserta una ecuación en el párrafo usando MathML (fallback a texto formateado).
    """
    if not LATEX_SUPPORT:
        run = paragraph.add_run(f" [{latex_code}] ")
        run.italic = True
        return

    try:
        # Limpiar el código LaTeX: quitar espacios extras y caracteres problemáticos
        latex_code = latex_code.strip()
        # Convertir a MathML
        mathml = latex2mathml(latex_code)
        # Extraer el contenido dentro de <math>...</math>
        root = ET.fromstring(mathml)
        # Buscar el elemento <math> (puede tener namespace)
        math_elem = root if root.tag.endswith('math') else root.find('.//{*}math')
        if math_elem is None:
            raise ValueError("No se encontró elemento math")
        # Crear elemento m:oMath en Word
        oMath = OxmlElement('m:oMath')
        # Copiar los hijos del math al oMath
        for child in math_elem:
            oMath.append(child)
        paragraph._p.append(oMath)
    except Exception as e:
        # Si falla, mostrar el LaTeX original con formato legible
        run = paragraph.add_run(f" {latex_code} ")
        run.italic = True
        run.font.color.rgb = None  # color normal

class BaseFormatter:
    def __init__(self, doc):
        self.doc = doc

    def _procesar_texto_con_ecuaciones(self, paragraph, texto):
        """
        Detecta fragmentos entre $...$ (inline) o $$...$$ (display)
        y los convierte a ecuaciones MathML o texto formateado.
        """
        # Patrón para ecuaciones en bloque $$ ... $$
        pattern_block = r'\$\$(.*?)\$\$'
        partes = re.split(pattern_block, texto, flags=re.DOTALL)
        for i, parte in enumerate(partes):
            if i % 2 == 1:  # Es ecuación en bloque
                add_math_equation(paragraph, parte.strip())
            else:
                # Ahora ecuaciones inline $ ... $
                inline_parts = re.split(r'\$(.*?)\$', parte)
                for j, inline in enumerate(inline_parts):
                    if j % 2 == 1:
                        add_math_equation(paragraph, inline.strip())
                    else:
                        if inline:
                            run = paragraph.add_run(inline)
                            run.font.size = Pt(11)
        return paragraph

    def agregar_parrafo_normal(self, texto):
        p = self.doc.add_paragraph()
        self._procesar_texto_con_ecuaciones(p, texto)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Cm(0.5)
        return p

    def agregar_titulo(self, texto, nivel=1):
        p = self.doc.add_paragraph()
        run = p.add_run(texto)
        run.bold = True
        if nivel == 1:
            run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        else:
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        return p

    def agregar_lista(self, texto):
        p = self.doc.add_paragraph()
        run = p.add_run("• ")
        run.font.size = Pt(11)
        self._procesar_texto_con_ecuaciones(p, texto)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(4)
        return p

    def procesar_linea(self, linea):
        raise NotImplementedError

class InglesFormatter(BaseFormatter):
    def procesar_linea(self, linea):
        linea = linea.strip()
        if not linea:
            return
        if linea.startswith('- ') or linea.startswith('• '):
            self.agregar_lista(linea[2:])
        else:
            self.agregar_parrafo_normal(linea)

class CiudadaniaFormatter(BaseFormatter):
    def procesar_linea(self, linea):
        linea = linea.strip()
        if not linea:
            return
        if re.match(r'^\d+\.\s+', linea) and not re.match(r'^\d+\.\d+\.', linea):
            self.agregar_titulo(linea, nivel=1)
        elif re.match(r'^\d+\.\d+\.\s+', linea):
            self.agregar_titulo(linea, nivel=2)
        elif linea.startswith('- ') or linea.startswith('• '):
            self.agregar_lista(linea[2:])
        else:
            self.agregar_parrafo_normal(linea)

class EstadisticaFormatter(BaseFormatter):
    def __init__(self, doc):
        super().__init__(doc)
        # Reemplazos para símbolos comunes (se aplican fuera de las ecuaciones LaTeX)
        self.reemplazos_texto = {
            r'\<=': '≤', r'\>=': '≥', r'\neq': '≠', r'\sqrt': '√',
            r'\alpha': 'α', r'\beta': 'β', r'\sigma': 'σ',
            r'\tilde': '~',  # simplificar tilde
            r'\mu': 'μ'
        }

    def _preconvertir_texto(self, texto):
        # Aplica reemplazos solo en texto que no está dentro de $...$ o $$...$$
        # Dividimos por ecuaciones para no tocarlas
        parts = re.split(r'(\$\$.*?\$\$|\$.*?\$)', texto, flags=re.DOTALL)
        resultado = []
        for part in parts:
            if part.startswith('$'):
                resultado.append(part)  # dejar las ecuaciones intactas
            else:
                for k, v in self.reemplazos_texto.items():
                    part = part.replace(k, v)
                # Convertir _digito a subíndice Unicode (solo si no es parte de LaTeX)
                def sub_repl(m):
                    return ''.join(chr(0x2080 + int(d)) for d in m.group(1))
                part = re.sub(r'_([0-9]+)', sub_repl, part)
                # ^digito a superíndice
                def sup_repl(m):
                    sup_map = {'0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵','6':'⁶','7':'⁷','8':'⁸','9':'⁹'}
                    return ''.join(sup_map.get(d,d) for d in m.group(1))
                part = re.sub(r'\^([0-9]+)', sup_repl, part)
                resultado.append(part)
        return ''.join(resultado)

    def procesar_linea(self, linea):
        linea = linea.strip()
        if not linea:
            return
        # Títulos ## y ###
        if linea.startswith('## '):
            self.agregar_titulo(linea[3:], nivel=1)
        elif linea.startswith('### '):
            self.agregar_titulo(linea[4:], nivel=2)
        elif linea.startswith('- '):
            texto_lista = linea[2:]
            # Aplicar conversión de símbolos (pero respetando ecuaciones)
            texto_lista = self._preconvertir_texto(texto_lista)
            self.agregar_lista(texto_lista)
        elif linea.startswith('> '):
            self.procesar_linea(linea[2:])
        else:
            texto = self._preconvertir_texto(linea)
            self.agregar_parrafo_normal(texto)

class GeneralFormatter(BaseFormatter):
    def procesar_linea(self, linea):
        linea = linea.strip()
        if not linea:
            return
        self.agregar_parrafo_normal(linea)

# Mapeo de tipos
FORMATEADORES = {
    'ingles': InglesFormatter,
    'ciudadania': CiudadaniaFormatter,
    'estadistica': EstadisticaFormatter,
    'general': GeneralFormatter,
}