# generadores/estadistica.py
import re
import os
import platform
from io import BytesIO

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .base import BaseGenerador

_COLORES = ['#2196F3','#FF5722','#4CAF50','#9C27B0',
            '#FF9800','#00BCD4','#E91E63','#795548']

def _parsear_chart_tag(tag_texto: str) -> dict | None:
    m = re.match(r'\[CHART:\s*(.*?)\]', tag_texto, re.IGNORECASE | re.DOTALL)
    if not m:
        return None
    contenido = m.group(1).strip()
    partes = [p.strip() for p in contenido.split('|')]
    if len(partes) < 4:
        return None
    try:
        tipo      = partes[0].lower()
        titulo    = partes[1]
        etiquetas = [e.strip() for e in partes[2].split(',')]
        valores   = [float(v.strip()) for v in partes[3].split(',')]
        if len(etiquetas) != len(valores):
            return None
        return {'tipo': tipo, 'titulo': titulo, 'etiquetas': etiquetas, 'valores': valores}
    except Exception:
        return None

def _generar_grafico_png(info: dict) -> BytesIO | None:
    tipo      = info['tipo']
    titulo    = info['titulo']
    etiquetas = info['etiquetas']
    valores   = info['valores']
    colores   = _COLORES[:len(valores)]

    try:
        fig, ax = plt.subplots(figsize=(7, 4.5))
        fig.patch.set_facecolor('white')

        if tipo in ('bar', 'column', 'columnas', 'barras'):
            bars = ax.bar(etiquetas, valores, color=colores, width=0.55,
                          edgecolor='white', linewidth=0.8)
            ax.set_ylabel('Valor', fontsize=10)
            ax.set_ylim(0, max(valores) * 1.25)
            for bar, val in zip(bars, valores):
                ax.text(bar.get_x() + bar.get_width()/2,
                        bar.get_height() + max(valores)*0.02,
                        f'{val:g}', ha='center', va='bottom', fontsize=9, fontweight='bold')
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            plt.xticks(rotation=15 if max(len(e) for e in etiquetas) > 8 else 0,
                       ha='right' if max(len(e) for e in etiquetas) > 8 else 'center', fontsize=9)

        elif tipo in ('barh', 'horizontal'):
            bars = ax.barh(etiquetas, valores, color=colores, height=0.55)
            ax.set_xlabel('Valor', fontsize=10)
            ax.set_xlim(0, max(valores) * 1.25)
            for bar, val in zip(bars, valores):
                ax.text(bar.get_width() + max(valores)*0.01, bar.get_y() + bar.get_height()/2,
                        f'{val:g}', va='center', fontsize=9, fontweight='bold')
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)

        elif tipo in ('pie', 'circulo', 'circular', 'torta'):
            wedges, texts, autotexts = ax.pie(
                valores, labels=etiquetas, colors=colores,
                autopct='%1.1f%%', startangle=90,
                pctdistance=0.75, labeldistance=1.1,
                wedgeprops=dict(width=0.6, edgecolor='white', linewidth=2)
            )
            for t in texts: t.set_fontsize(9)
            for at in autotexts: at.set_fontsize(9); at.set_fontweight('bold')
            ax.axis('equal')

        elif tipo in ('line', 'linea', 'lineas'):
            ax.plot(etiquetas, valores, marker='o', color=colores[0],
                    linewidth=2.5, markersize=7, markerfacecolor='white',
                    markeredgewidth=2.5)
            ax.fill_between(range(len(etiquetas)), valores,
                            alpha=0.12, color=colores[0])
            ax.set_ylim(min(valores)*0.85, max(valores)*1.2)
            for i, (et, val) in enumerate(zip(etiquetas, valores)):
                ax.text(i, val + max(valores)*0.02, f'{val:g}',
                        ha='center', fontsize=9, fontweight='bold')
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)

        else:
            ax.bar(etiquetas, valores, color=colores)

        ax.set_title(titulo, fontsize=13, fontweight='bold', pad=12)
        plt.tight_layout()

        buf = BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                    facecolor='white')
        buf.seek(0)
        plt.close(fig)
        return buf

    except Exception as e:
        print(f"[CHART ERROR] {e}")
        plt.close('all')
        return None

def _set_cell_bg(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


class GeneradorEstadistica(BaseGenerador):
    
    def __init__(self, nombre_archivo="Tarea_Final.docx"):
        super().__init__(nombre_archivo)
        self.ruta_absoluta = os.path.abspath(nombre_archivo)
        self._tabla_buffer = []
        self._dentro_tabla = False
        self._configurar_estilos_doc()

    def _configurar_estilos_doc(self):
        estilo = self.doc.styles['Normal']
        estilo.font.name = "Arial"
        estilo.font.size = Pt(11)
        estilo.paragraph_format.line_spacing = 1.5

    def _insertar_grafico(self, tag_texto: str):
        info = _parsear_chart_tag(tag_texto)
        if not info:
            p = self.doc.add_paragraph()
            p.add_run(f"[Gráfico no pudo generarse: {tag_texto[:80]}]").italic = True
            return
        buf = _generar_grafico_png(info)
        if buf is None:
            p = self.doc.add_paragraph()
            p.add_run(f"[Error generando gráfico: {info['titulo']}]").italic = True
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        run.add_picture(buf, width=Inches(5.5))
        caption = self.doc.add_paragraph()
        caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_c = caption.add_run(f"Figura: {info['titulo']}")
        run_c.italic = True
        run_c.font.size = Pt(9)
        run_c.font.name = "Arial"
        caption.paragraph_format.space_after = Pt(10)

    def _es_linea_tabla(self, linea: str) -> bool:
        s = linea.strip()
        return s.startswith('|') and s.endswith('|') and '|' in s[1:-1]

    def _plasmar_tabla_markdown(self, lineas: list):
        filas_datos = []
        for linea in lineas:
            linea = linea.strip()
            if not linea:
                continue
            interior = linea.strip('|')
            if all(re.match(r'^[\s\-:]+$', c.strip()) for c in interior.split('|') if c.strip()):
                continue
            celdas = [c.strip() for c in interior.split('|')]
            filas_datos.append(celdas)

        if not filas_datos:
            return

        num_cols = max(len(f) for f in filas_datos)
        for fila in filas_datos:
            while len(fila) < num_cols:
                fila.append('')

        tabla = self.doc.add_table(rows=len(filas_datos), cols=num_cols)
        tabla.style = 'Table Grid'

        for i, fila in enumerate(filas_datos):
            for j, celda_txt in enumerate(fila):
                cell = tabla.cell(i, j)
                if i == 0:
                    _set_cell_bg(cell, '2E74B5')
                    p = cell.paragraphs[0]
                    run = p.add_run(celda_txt)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.name = "Arial"
                    run.font.size = Pt(11)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                else:
                    if i % 2 == 0:
                        _set_cell_bg(cell, 'EBF3FB')
                    p = cell.paragraphs[0]
                    self._procesar_texto_con_ecuaciones(p, celda_txt)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def agregar_titulo(self, texto, nivel=1):
        p = self.doc.add_paragraph()
        run = p.add_run(texto.strip())
        run.bold = True
        run.font.name = "Arial"
        if nivel == 1:
            run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
        elif nivel == 2:
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        else:
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.line_spacing = 1.5
        return p

    def agregar_parrafo(self, texto, sangria=True):
        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        self._procesar_texto_con_ecuaciones(p, texto)
        return p

    def agregar_lista(self, texto, nivel=0):
        p = self.doc.add_paragraph()
        run = p.add_run(f"{'  ' * nivel}• ")
        run.font.size = Pt(11)
        run.font.name = "Arial"
        p.paragraph_format.left_indent = Cm(0.8 + nivel * 0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        self._procesar_texto_con_ecuaciones(p, texto)
        return p

    def agregar_ecuacion_bloque(self, latex_code: str):
        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        self._insertar_ecuacion(p, latex_code)
        return p

    def procesar_linea(self, linea):
        linea_raw = linea.rstrip()

        if self._es_linea_tabla(linea_raw):
            self._dentro_tabla = True
            self._tabla_buffer.append(linea_raw)
            return

        if self._dentro_tabla:
            self._plasmar_tabla_markdown(self._tabla_buffer)
            self._tabla_buffer = []
            self._dentro_tabla = False

        linea = linea_raw.strip()
        if not linea:
            return

        if re.match(r'\[CHART:', linea, re.IGNORECASE):
            self._insertar_grafico(linea)
            return

        m_eq_sola = re.match(r'^\s*\[EQ:(.*?)\]\s*$', linea, re.DOTALL)
        if m_eq_sola:
            self.agregar_ecuacion_bloque(m_eq_sola.group(1).strip())
            return

        if linea.startswith('#### '):
            self.agregar_titulo(linea[5:], nivel=3)
        elif linea.startswith('### '):
            self.agregar_titulo(linea[4:], nivel=3)
        elif linea.startswith('## '):
            self.agregar_titulo(linea[3:], nivel=2)
        elif linea.startswith('# '):
            self.agregar_titulo(linea[2:], nivel=1)
        elif re.match(r'^\d+\.\d+\.\d+\.\s+', linea):
            self.agregar_titulo(linea, nivel=3)
        elif re.match(r'^\d+\.\d+\.\s+', linea):
            self.agregar_titulo(linea, nivel=2)
        elif re.match(r'^\d+\.\s+[A-ZÁÉÍÓÚÜÑ\w]', linea) and len(linea) < 120:
            posible_titulo = re.sub(r'^\d+\.\s+', '', linea)
            if len(posible_titulo.split()) <= 12:
                self.agregar_titulo(linea, nivel=1)
            else:
                self.agregar_parrafo(linea)

        elif linea.startswith('- ') or linea.startswith('• ') or linea.startswith('* '):
            self.agregar_lista(linea[2:])
        elif re.match(r'^\d+\)\s+', linea):
            self.agregar_lista(re.sub(r'^\d+\)\s+', '', linea))

        else:
            self.agregar_parrafo(linea)

    def procesar_todo(self, texto_completo):
        lineas = texto_completo.split('\n')
        for linea in lineas:
            self.procesar_linea(linea)
        if self._dentro_tabla and self._tabla_buffer:
            self._plasmar_tabla_markdown(self._tabla_buffer)
            self._tabla_buffer = []
            self._dentro_tabla = False

    def guardar(self):
        self.doc.save(self.ruta_absoluta)
        if platform.system() == "Windows":
            self._buildup_windows()

    def _buildup_windows(self):
        if platform.system() == "Windows":
            try:
                import win32com.client
                import pythoncom
                pythoncom.CoInitialize()
                app = win32com.client.Dispatch("Word.Application")
                app.Visible = False
                doc = app.Documents.Open(self.ruta_absoluta)
                doc.OMaths.BuildUp()
                doc.Save()
                doc.Close()
                app.Quit()
                pythoncom.CoUninitialize()
            except Exception as e:
                print(f"[BuildUp Windows] {e}")

    def optimizar_y_guardar_perfecto(self):
        self.guardar()