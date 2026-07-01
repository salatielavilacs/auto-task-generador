# generadores/base.py
import re
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml

try:
    from latex2mathml.converter import convert as latex2mathml
    LATEX_SUPPORT = True
except ImportError:
    LATEX_SUPPORT = False
    latex2mathml = None

def convert_mathml_to_omml_str(elem) -> str:
    """Convierte de forma recursiva los elementos MathML a la estructura nativa OMML."""
    tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
    
    if tag in ('math', 'mrow', 'style', 'semantics', 'mstyle'):
        return "".join(convert_mathml_to_omml_str(child) for child in elem)
        
    elif tag in ('mi', 'mn', 'mo', 'mtext'):
        val = elem.text or ""
        val = val.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        return f"<m:r><m:t>{val}</m:t></m:r>"
        
    elif tag == 'mfrac':
        if len(elem) < 2:
            return ""
        num = convert_mathml_to_omml_str(elem[0])
        den = convert_mathml_to_omml_str(elem[1])
        return (
            "<m:f>"
            "<m:fPr><m:type m:val=\"bar\"/></m:fPr>"
            f"<m:num>{num}</m:num>"
            f"<m:den>{den}</m:den>"
            "</m:f>"
        )
    # ─── NUEVO: SOPORTE NATIVO PARA MATRICES EN OMML ───
    elif tag == 'mtable':
        inner = "".join(convert_mathml_to_omml_str(child) for child in elem)
        return (
            "<m:m>"
            "<m:mPr>"
            "<m:baseJc m:val=\"center\"/>"  # Alineación vertical centrada
            "<m:plcHide m:val=\"1\"/>"
            "</m:mPr>"
            f"{inner}"
            "</m:m>"
        )
        
    elif tag == 'mtr':
        inner = "".join(convert_mathml_to_omml_str(child) for child in elem)
        return f"<m:mr>{inner}</m:mr>"
        
    elif tag == 'mtd':
        inner = "".join(convert_mathml_to_omml_str(child) for child in elem)
        return f"<m:e>{inner}</m:e>"
    # ──────────────────────────────────────────────────    
    elif tag == 'msup':
        if len(elem) < 2:
            return ""
        base = convert_mathml_to_omml_str(elem[0])
        sup = convert_mathml_to_omml_str(elem[1])
        return (
            "<m:sSup>"
            "<m:sSupPr/>"
            f"<m:e>{base}</m:e>"
            f"<m:sup>{sup}</m:sup>"
            "</m:sSup>"
        )
        
    elif tag == 'msub':
        if len(elem) < 2:
            return ""
        base = convert_mathml_to_omml_str(elem[0])
        sub = convert_mathml_to_omml_str(elem[1])
        return (
            "<m:sSub>"
            "<m:sSubPr/>"
            f"<m:e>{base}</m:e>"
            f"<m:sub>{sub}</m:sub>"
            "</m:sSub>"
        )
        
    elif tag == 'msubsup':
        if len(elem) < 3:
            return ""
        base = convert_mathml_to_omml_str(elem[0])
        sub = convert_mathml_to_omml_str(elem[1])
        sup = convert_mathml_to_omml_str(elem[2])
        return (
            "<m:sSubSup>"
            "<m:sSubSupPr/>"
            f"<m:e>{base}</m:e>"
            f"<m:sub>{sub}</m:sub>"
            f"<m:sup>{sup}</m:sup>"
            "</m:sSubSup>"
        )
        
    elif tag == 'msqrt':
        inner = "".join(convert_mathml_to_omml_str(child) for child in elem)
        return (
            "<m:rad>"
            "<m:radPr><m:degHide m:val=\"1\"/></m:radPr>"
            "<m:deg/>"
            f"<m:e>{inner}</m:e>"
            "</m:rad>"
        )
        
    elif tag == 'mroot':
        if len(elem) < 2:
            return ""
        base = convert_mathml_to_omml_str(elem[0])
        deg = convert_mathml_to_omml_str(elem[1])
        return (
            "<m:rad>"
            "<m:radPr/>"
            f"<m:deg>{deg}</m:deg>"
            f"<m:e>{base}</m:e>"
            "</m:rad>"
        )
        
    elif tag in ('munder', 'mover', 'munderover'):
        if tag == 'munder':
            if len(elem) < 2: return ""
            base = convert_mathml_to_omml_str(elem[0])
            sub = convert_mathml_to_omml_str(elem[1])
            return f"<m:limLow><m:limLowPr/><m:e>{base}</m:e><m:lim>{sub}</m:lim></m:limLow>"
        elif tag == 'mover':
            if len(elem) < 2: return ""
            base = convert_mathml_to_omml_str(elem[0])
            sup = convert_mathml_to_omml_str(elem[1])
            return f"<m:limUpp><m:limUppPr/><m:e>{base}</m:e><m:lim>{sup}</m:lim></m:limUpp>"
        elif tag == 'munderover':
            if len(elem) < 3: return ""
            base = convert_mathml_to_omml_str(elem[0])
            sub = convert_mathml_to_omml_str(elem[1])
            sup = convert_mathml_to_omml_str(elem[2])
            
            is_nary = False
            if len(elem[0].text or "") > 0:
                char = elem[0].text.strip()
                if char in ("∑", "∫", "∏", "⋃", "⋂", "\\sum", "\\int", "\\prod"):
                    is_nary = True
            
            if is_nary:
                char = elem[0].text.strip()
                return (
                    "<m:nary>"
                    f"<m:naryPr><m:chr m:val=\"{char}\"/><m:limLoc m:val=\"undOvr\"/><m:subHide m:val=\"0\"/><m:supHide m:val=\"0\"/></m:naryPr>"
                    f"<m:sub>{sub}</m:sub>"
                    f"<m:sup>{sup}</m:sup>"
                    "<m:e/>"
                    "</m:nary>"
                )
            else:
                return (
                    "<m:limUpp><m:limUppPr/>"
                    f"<m:e><m:limLow><m:limLowPr/><m:e>{base}</m:e><m:lim>{sub}</m:lim></m:limLow></m:e>"
                    f"<m:lim>{sup}</m:lim>"
                    "</m:limUpp>"
                )
    
    return "".join(convert_mathml_to_omml_str(child) for child in elem)


# Detecta notación de potencia tipo "base^exponente" en texto plano, p.ej.:
#   (37 - 33.70)^2 / 33.70 = 0.32   ->   (37 - 33.70)²  con el 2 en superíndice real
# La base puede ser un paréntesis "(...)" o un token alfanumérico; el exponente
# puede ser un número (negativo o decimal), un token, o ir entre paréntesis.
_RE_POTENCIA = re.compile(
    r'(\([^()]+\)|[A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9._]+)\^(-?[A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9.]+|\([^()]+\))'
)


class BaseGenerador:
    def __init__(self, nombre_archivo="Tarea_Final.docx"):
        self.doc = Document()
        self.nombre_archivo = nombre_archivo

    def _agregar_texto_con_potencias(self, paragraph, texto, bold=False,
                                      font_name=None, font_size=None):
        """Agrega texto al párrafo convirtiendo automáticamente cualquier
        notación 'base^exponente' en superíndice nativo de Word.

        Esto es independiente del sistema operativo (no usa win32com ni
        depende de que la IA haya envuelto la expresión en [EQ: ...]):
        funciona igual en Linux, macOS y Windows porque python-docx maneja
        el superíndice como un simple atributo de formato del run.
        """
        if not texto:
            return

        def _aplicar_formato(run, superindice=False):
            run.bold = bold
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = font_size
            if superindice:
                run.font.superscript = True

        pos = 0
        for m in _RE_POTENCIA.finditer(texto):
            if m.start() > pos:
                _aplicar_formato(paragraph.add_run(texto[pos:m.start()]))
            base, exp = m.group(1), m.group(2)
            _aplicar_formato(paragraph.add_run(base))
            _aplicar_formato(paragraph.add_run(exp.strip('()')), superindice=True)
            pos = m.end()
        if pos < len(texto):
            _aplicar_formato(paragraph.add_run(texto[pos:]))

    def _insertar_ecuacion(self, paragraph, latex_code):
        """Convierte código LaTeX a un objeto OMML nativo de Word y lo inserta en el párrafo."""
        latex_code = latex_code.strip()
        if LATEX_SUPPORT and latex2mathml:
            try:
                mathml_str = latex2mathml(latex_code)
                root = ET.fromstring(mathml_str)
                omml_content = convert_mathml_to_omml_str(root)
                omml_xml = f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{omml_content}</m:oMath>'
                element = parse_xml(omml_xml)
                paragraph._p.append(element)
                return
            except Exception as e:
                print(f"[Math Engine Warning] {e}. Ejecutando fallback a texto plano.")
        
        # Fallback seguro en caso de error o ausencia de la librería
        fallback_xml = f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:r><m:t>{latex_code}</m:t></m:r></m:oMath>'
        paragraph._p.append(parse_xml(fallback_xml))

    def _procesar_texto_con_ecuaciones(self, paragraph, texto):
        """Busca [EQ: ...] y los convierte en ecuaciones; además aplica negritas **texto**."""
        texto_normalizado = re.sub(r'(\\\w+\s*(?:\\leq|>|=)\\geq\s*\d+)', r'[EQ:\1]', texto)
        partes_eq = re.split(r'(\[EQ:.*?\])', texto_normalizado)
        for parte in partes_eq:
            if parte.startswith('[EQ:') and parte.endswith(']'):
                codigo = parte[4:-1].strip()
                self._insertar_ecuacion(paragraph, codigo)
            else:
                sub_partes = re.split(r'(\*\*.*?\*\*)', parte)
                for sub in sub_partes:
                    if sub.startswith('**') and sub.endswith('**'):
                        self._agregar_texto_con_potencias(paragraph, sub[2:-2], bold=True)
                    else:
                        if sub:
                            self._agregar_texto_con_potencias(paragraph, sub)

    def agregar_parrafo(self, texto, sangria=True):
        p = self.doc.add_paragraph()
        self._procesar_texto_con_ecuaciones(p, texto)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        return p

    def agregar_titulo(self, texto, nivel=1):
        p = self.doc.add_paragraph()
        run = p.add_run(texto)
        run.bold = True
        if nivel == 1:
            run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        else:
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        return p

    def agregar_lista(self, texto, nivel=0):
        p = self.doc.add_paragraph()
        run = p.add_run(f"{'  ' * nivel}• {texto}")
        run.font.size = Pt(11)
        p.paragraph_format.left_indent = Cm(0.8 + nivel * 0.5)
        p.paragraph_format.space_after = Pt(4)
        return p

    def procesar_linea(self, linea):
        raise NotImplementedError

    def agregar_texto(self, texto):
        for linea in texto.split('\n'):
            self.procesar_linea(linea)

    def guardar(self):
        self.doc.save(self.nombre_archivo)