# generadores/general.py
import re
import os
import platform
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from .base import BaseGenerador

class GeneradorEstadistica(BaseGenerador):
    
    def __init__(self, nombre_archivo="Tarea_Final.docx"):
        super().__init__(nombre_archivo)
        self.ruta_absoluta = os.path.abspath(nombre_archivo)

    def procesar_todo(self, texto_completo):
        lineas = texto_completo.split('\n')
        for linea in lineas:
            if not linea.strip():
                continue
            
            if linea.strip().startswith("#"):
                self.agregar_titulo(linea.replace("#", "").strip())
            else:
                p = self.doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                self._procesar_texto_con_ecuaciones(p, linea)

    def agregar_parrafo(self, texto, sangria=True):
        p = self.doc.add_paragraph()
        self._procesar_texto_con_ecuaciones(p, texto)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        return p

    def agregar_lista(self, texto, nivel=0):
        p = self.doc.add_paragraph()
        run = p.add_run(f"{'  ' * nivel}• ")
        run.font.size = Pt(11)
        self._procesar_texto_con_ecuaciones(p, texto)
        p.paragraph_format.left_indent = Cm(0.8 + nivel * 0.5)
        p.paragraph_format.space_after = Pt(4)
        return p

    def procesar_linea(self, linea):
        linea = linea.strip()
        if not linea:
            return
        if linea.startswith('## '):
            self.agregar_titulo(linea[3:], nivel=1)
        elif linea.startswith('### '):
            self.agregar_titulo(linea[4:], nivel=2)
        elif linea.startswith('- '):
            self.agregar_lista(linea[2:])
        else:
            self.agregar_parrafo(linea)

    def optimizar_y_guardar_perfecto(self):
        self.doc.save(self.ruta_absoluta)
        
        if platform.system() == "Windows":
            try:
                import win32com.client
                word_app = win32com.client.Dispatch("Word.Application")
                word_app.Visible = False 
                doc_win = word_app.Documents.Open(self.ruta_absoluta)
                word_app.Selection.WholeStory()
                doc_win.OMaths.BuildUp()
                doc_win.Save()
                doc_win.Close()
                word_app.Quit()
                print(f"Éxito en la optimización sobre plataforma Windows.")
            except Exception as e:
                print(f"Soporte Word.Application omitido: {e}")